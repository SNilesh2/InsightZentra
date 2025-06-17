import os
import threading
from io import BytesIO
from flask import Flask, render_template, request, redirect, url_for, send_file, make_response 
from flask import jsonify
from gtts import gTTS
from deep_translator import GoogleTranslator
from fpdf import FPDF
from docx import Document
from langchain_groq import ChatGroq
from langchain_community.document_loaders import WebBaseLoader
from langchain_huggingface import HuggingFaceEmbeddings
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.chains.combine_documents import create_stuff_documents_chain
from langchain_core.prompts import ChatPromptTemplate
from langchain.chains.question_answering import load_qa_chain
from langchain.chains import create_retrieval_chain
from langchain_community.vectorstores import FAISS
from dotenv import load_dotenv
load_dotenv()

groq_api_key = os.getenv("GROQ_API_KEY")
os.environ["USER_AGENT"] = "Mozilla/5.0 (Windows NT 10.0; Win64; x64)"

# Language code map for gTTS
LANGUAGE_CODE_MAP = {
    "en": "en",           # English
    "zh-CN": "zh-cn",     # Mandarin Chinese
    "hi": "hi",           # Hindi
    "es": "es",           # Spanish
    "fr": "fr",           # French
    "ar": "ar",           # Modern Standard Arabic
    "bn": "bn",           # Bengali
    "ru": "ru",           # Russian
    "ta": "ta",           # Tamil
    "kn": "kn",           # Kannada
    "ml": "ml",           # Malayalam
    "te": "te"            # Telugu
}

def initialize_session_state():
    session_state['embeddings'] = HuggingFaceEmbeddings(model_name="sentence-transformers/all-MiniLM-L6-v2")
    session_state['qna'] = []
    session_state['current_url'] = None
    session_state['summary_language'] = 'en'
    session_state['qna_language'] = 'en'
    session_state['url'] = ""
    session_state['summary'] = ""
    session_state['prerequisites'] = []
    session_state['history'] = []

session_state = {}
initialize_session_state()

llm = ChatGroq(groq_api_key=groq_api_key, model_name="llama3-8b-8192")


prompt_template = ChatPromptTemplate.from_template(
    """
    When input is a question, provide a concise answer that directly addresses the main point without adding unrelated context or excessive detail.
    - For general or conceptual questions (e.g., "What is a graph?"), limit the answer to a 1-2 sentence definition or explanation that is simple and straightforward.
    - For technical or detailed questions, provide a focused answer in 2-3 sentences, addressing only the main technical aspects or requested details.

    For a request to summarize, focus strictly on the provided article URL’s main ideas, core concepts, and examples. 
    - Exclude unrelated details, references to other articles, or platform-specific information.
    - Start with a brief overview of the primary topic, followed by 2-3 key points or concepts discussed in the article.
    - Summarize in 3-5 sentences to capture the essential takeaways.

    Ensure all responses are:
    1. Accurate and technically correct.
    2. Tailored to the complexity of the question or request.
    3. Clear, concise, and aligned with the user's input.
    4. Structured appropriately for readability:
       - Use bullet points for key findings where applicable.
       - Avoid overloading the user with unnecessary information.

    <context>{context}</context>
    """
)

app = Flask(__name__)
terminate_server = False

# Directory to store audio files
AUDIO_DIR = "static/audio"

if not os.path.exists(AUDIO_DIR):
    os.makedirs(AUDIO_DIR)

@app.route('/')
def home():
    return render_template('index.html', 
                           qna=session_state.get('qna', []), 
                           summary_language=session_state.get('summary_language', 'en'), 
                           qna_language=session_state.get('qna_language', 'en'), 
                           url=session_state.get('url', ''), 
                           summary=session_state.get('summary', ''), 
                           prerequisites=session_state.get('prerequisites', []),
                           history=session_state.get('history', []),
                           summary_audio_exists=os.path.exists(f"{AUDIO_DIR}/summary_audio.mp3"))


@app.route('/translate_all', methods=['POST'])
def translate_all():
    selected_language = request.form['language']
    session_state['summary_language'] = selected_language
    session_state['qna_language'] = selected_language

    response_data = {'success': False}

    try:
         # Translate Summary from original English summary
        if 'original_summary' in session_state and session_state['original_summary']:
            translated_summary = GoogleTranslator(source='en', target=selected_language).translate(session_state['original_summary'])
            session_state['summary'] = translated_summary  # Store the translated version
            generate_audio(translated_summary, "summary_audio.mp3", selected_language)
            response_data['summary'] = translated_summary

        # Translate Q&A Answers
        translated_qna = []
        if 'qna' in session_state:
            for idx, qna in enumerate(session_state['qna']):
                translated_answer = GoogleTranslator(source='en', target=selected_language).translate(qna['original_answer'])
                session_state['qna'][idx]['translated_answer'] = translated_answer
                session_state['qna'][idx]['answer'] = translated_answer  # for display
                generate_audio(translated_answer, f"answer_{idx + 1}.mp3", selected_language)
                translated_qna.append({'question': qna['question'], 'answer': translated_answer})

        response_data['qna'] = translated_qna
        response_data['success'] = True

    except Exception as e:
        response_data['error'] = str(e)

    return jsonify(response_data)



@app.route('/submit_url', methods=['POST'])
def submit_url():
    url = request.form['url'].strip()
    if not url:
        return render_template('index.html', error="Please enter a valid URL.",
                               qna=session_state.get('qna', []),
                               summary_language=session_state.get('summary_language', 'en'),
                               qna_language=session_state.get('qna_language', 'en'),
                               url=session_state.get('url', ''),
                               summary=session_state.get('summary', ''),
                               prerequisites=session_state.get('prerequisites', []),
                               history=session_state.get('history', []),
                               summary_audio_exists=os.path.exists(f"{AUDIO_DIR}/summary_audio.mp3"))

    if not url.startswith("http://") and not url.startswith("https://"):
        url = "https://" + url

    try:
        if not url.startswith("https://") or not "." in url.split("//")[1]:
            return render_template('index.html', error="Invalid URL. Please enter a valid web URL.",
                                   qna=session_state.get('qna', []),
                                   summary_language=session_state.get('summary_language', 'en'),
                                   qna_language=session_state.get('qna_language', 'en'),
                                   url=session_state.get('url', ''),
                                   summary=session_state.get('summary', ''),
                                   prerequisites=session_state.get('prerequisites', []),
                                   history=session_state.get('history', []),
                                   summary_audio_exists=os.path.exists(f"{AUDIO_DIR}/summary_audio.mp3"))

        session_state['url'] = url

        if 'loader' not in session_state or session_state.get('current_url') != url:
            session_state['loader'] = WebBaseLoader(url)
            session_state['docs'] = session_state['loader'].load()
            session_state['text_splitter'] = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=200)
            session_state['final_documents'] = session_state['text_splitter'].split_documents(session_state['docs'])
            session_state['vectors'] = FAISS.from_documents(session_state['final_documents'], session_state['embeddings'])
            session_state['current_url'] = url

        # Clear Q&A on new URL
        session_state['qna'] = []

        # New summary-only prompt to ensure clean summary without Q&A text
        prompt_template_summary_only = ChatPromptTemplate.from_template(
            """
            You are asked to provide only a concise summary of the article content.
            - Focus strictly on main ideas, key points, and core concepts.
            - Provide a clean summary of 3-5 sentences.
            - Do NOT include any questions or answers.

            <context>{context}</context>
            """
        )

        document_chain = create_stuff_documents_chain(llm, prompt_template_summary_only)
        retriever = session_state['vectors'].as_retriever()
        retrieval_chain = create_retrieval_chain(retriever, document_chain)

        response = retrieval_chain.invoke({
            "input": f"Summarize the article from the URL: {url}."
        })

        # Save original summary in English for translation usage
        session_state['original_summary'] = response['answer']
        session_state['summary'] = response['answer']

        # Generate audio for the summary
        generate_audio(session_state['summary'], "summary_audio.mp3", session_state['summary_language'])

    except Exception as e:
        session_state['summary'] = f"An error occurred during summarization: {str(e)}"
        session_state['prerequisites'] = []

    return redirect(url_for('home'))



@app.route('/get_answer', methods=['POST'])
def get_answer():
    user_prompt = request.form['prompt']
    if not user_prompt.strip():
        return redirect(url_for('home'))

    try:
        # Define prompt template using context and question placeholders
        prompt_template = ChatPromptTemplate.from_template("""
        You are a helpful assistant. Answer the user's question using ONLY the context provided.
        Be concise and accurate. Do not add any summaries or unrelated information.

        <context>
        {context}
        </context>

        Question: {question}
        Answer:
        """)

        # Build the QA chain without unsupported arguments
        qa_chain = load_qa_chain(llm, chain_type="stuff", prompt=prompt_template)

        # Get retriever from session state
        retriever = session_state['vectors'].as_retriever()

        # Retrieve relevant documents
        relevant_docs = retriever.get_relevant_documents(user_prompt)

        # Run the QA chain with documents and question
        answer = qa_chain.run(input_documents=relevant_docs, question=user_prompt).strip()

        if not answer:
            answer = "Sorry, I couldn't find an answer based on the article."

        # Store the Q&A entry
        qna_entry = {
            'question': user_prompt,
            'answer': answer,
            'original_answer': answer
        }

        if 'qna' not in session_state:
            session_state['qna'] = []

        session_state['qna'].append(qna_entry)

        # Generate audio for answer
        generate_audio(answer, f"answer_{len(session_state['qna'])}.mp3", session_state.get('qna_language', 'en'))

    except Exception as e:
        qna_entry = {
            'question': user_prompt,
            'answer': f"An error occurred: {str(e)}",
            'original_answer': f"An error occurred: {str(e)}"
        }

        if 'qna' not in session_state:
            session_state['qna'] = []

        session_state['qna'].append(qna_entry)

    return redirect(url_for('home'))







@app.route('/audio/<filename>')
def get_audio(filename):
    """Serve the audio file if it exists."""
    file_path = os.path.join(AUDIO_DIR, filename)
    if os.path.exists(file_path):
        return send_file(file_path)
    else:
        return "Audio not available", 404

@app.route('/export', methods=['POST'])
def export():
    export_type = request.form['export_type']
    filename = f"export.{export_type}"

    if export_type == 'pdf':
        pdf = FPDF()
        pdf.add_page()

        # Add FreeSerif regular and bold fonts (make sure these .ttf files exist in fonts/)
        pdf.add_font('FreeSerif', '', 'fonts/FreeSerif.ttf', uni=True)
        pdf.add_font('FreeSerif', 'B', 'fonts/FreeSerifBold.ttf', uni=True)

        pdf.set_font("FreeSerif", size=14)
        pdf.cell(0, 10, "Article Summary and Q&A", ln=True, align='C')
        pdf.ln(10)

        pdf.set_font("FreeSerif", style='B', size=12)
        pdf.cell(0, 10, "Summary:", ln=True)
        pdf.set_font("FreeSerif", size=12)

        summary_text = session_state.get('translated_summary') or session_state.get('summary') or "No summary available"
        pdf.multi_cell(0, 10, summary_text)
        pdf.ln(10)

        pdf.set_font("FreeSerif", style='B', size=12)
        pdf.cell(0, 10, "Questions and Answers:", ln=True)
        pdf.set_font("FreeSerif", size=12)

        for qna in session_state.get('qna', []):
            question = qna.get('question', 'No question')
            answer = qna.get('translated_answer') or qna.get('answer') or "No answer available"
            pdf.cell(0, 10, f"Q: {question}", ln=True)
            pdf.multi_cell(0, 10, f"A: {answer}")
            pdf.ln(5)

        response = make_response(pdf.output(dest='S').encode('latin1', 'replace'))
        response.headers['Content-Type'] = 'application/pdf'
        response.headers['Content-Disposition'] = f'attachment; filename="{filename}"'
        return response

    elif export_type == 'docx':
        doc = Document()
        doc.add_heading('Article Summary and Q&A', level=1)

        doc.add_heading('Summary', level=2)
        summary_text = session_state.get('translated_summary') or session_state.get('summary') or "No summary available"
        doc.add_paragraph(summary_text)

        doc.add_heading('Questions and Answers', level=2)
        for qna in session_state.get('qna', []):
            question = qna.get('question', 'No question')
            answer = qna.get('translated_answer') or qna.get('answer') or "No answer available"
            doc.add_paragraph(f"Q: {question}", style='List Bullet')
            doc.add_paragraph(f"A: {answer}", style='Body Text')

        doc_io = BytesIO()
        doc.save(doc_io)
        doc_io.seek(0)
        return send_file(doc_io,
                         mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                         as_attachment=True, download_name=filename)

    return "Unsupported format", 400



@app.route('/new', methods=['POST'])
def new():
    # Save current session to history
    if session_state['url']:
        session_state['history'].append({
            'url': session_state['url'],
            'summary': session_state['summary'],
            'qna': session_state['qna']
        })

    # Clear current session
    session_state['qna'] = []
    session_state['url'] = ""
    session_state['summary'] = ""
    session_state['prerequisites'] = []
    session_state['current_url'] = None
    return redirect(url_for('home'))

@app.route('/exit', methods=['POST'])
def exit_app():
    global terminate_server
    terminate_server = True
    threading.Thread(target=shutdown_server).start()
    return "The application is shutting down..."

def shutdown_server():
    if terminate_server:
        func = request.environ.get('werkzeug.server.shutdown')
        if func:
            func()
        else:
            os._exit(0)

def generate_audio(text, filename, language):
    """Generate audio file from text using gTTS."""
    tts_language = LANGUAGE_CODE_MAP.get(language, "en")  # Default to English if language not supported
    
    try:
        if text:
            tts = gTTS(text=text, lang=tts_language)
            tts.save(os.path.join(AUDIO_DIR, filename))
    except Exception as e:
        print(f"Error generating audio for language '{language}': {e}")

if __name__ == '__main__':
    app.run(debug=True, use_reloader=False)